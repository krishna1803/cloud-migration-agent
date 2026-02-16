"""
Document processing utilities for the Cloud Migration Agent.

Handles extraction of text and data from PDF, DOCX, and Excel (BOM) files.
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument

from src.utils.logger import logger


class DocumentProcessor:
    """Process various document types for migration evidence"""
    
    @staticmethod
    def process_pdf(file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF document.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            reader = PdfReader(file_path)
            
            text_content = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                text_content.append({
                    "page": page_num,
                    "text": page_text
                })
            
            result = {
                "file_path": file_path,
                "file_type": "pdf",
                "num_pages": len(reader.pages),
                "content": text_content,
                "full_text": "\n".join([p["text"] for p in text_content]),
                "metadata": reader.metadata if reader.metadata else {}
            }
            
            logger.info(f"Processed PDF: {file_path} ({len(reader.pages)} pages)")
            return result
            
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def process_docx(file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = DocxDocument(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            result = {
                "file_path": file_path,
                "file_type": "docx",
                "num_paragraphs": len(paragraphs),
                "num_tables": len(tables),
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": "\n".join(paragraphs),
                "metadata": {
                    "author": doc.core_properties.author,
                    "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None
                }
            }
            
            logger.info(
                f"Processed DOCX: {file_path} "
                f"({len(paragraphs)} paragraphs, {len(tables)} tables)"
            )
            return result
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def process_excel_bom(file_path: str) -> Dict[str, Any]:
        """
        Process Bill of Materials (BOM) from Excel file.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with parsed BOM data
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to list of dictionaries
                records = df.to_dict('records')
                
                sheets_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "num_rows": len(df),
                    "data": records
                }
            
            # Try to identify resource and cost information
            resources = []
            total_cost = 0.0
            
            for sheet_name, sheet_data in sheets_data.items():
                for record in sheet_data["data"]:
                    # Look for common BOM fields
                    resource = {}
                    
                    # Try different column name variations
                    for key in record.keys():
                        key_lower = str(key).lower()
                        
                        if any(term in key_lower for term in ['service', 'resource', 'component']):
                            resource['name'] = record[key]
                        elif any(term in key_lower for term in ['type', 'category']):
                            resource['type'] = record[key]
                        elif any(term in key_lower for term in ['instance', 'size', 'sku']):
                            resource['instance_type'] = record[key]
                        elif any(term in key_lower for term in ['quantity', 'count']):
                            resource['quantity'] = record[key]
                        elif any(term in key_lower for term in ['cost', 'price', 'amount']):
                            try:
                                cost = float(str(record[key]).replace('$', '').replace(',', ''))
                                resource['monthly_cost'] = cost
                                total_cost += cost
                            except (ValueError, AttributeError):
                                pass
                        elif any(term in key_lower for term in ['region', 'location']):
                            resource['region'] = record[key]
                    
                    if resource and resource.get('name'):
                        resources.append(resource)
            
            result = {
                "file_path": file_path,
                "file_type": "excel_bom",
                "sheets": sheets_data,
                "resources": resources,
                "total_monthly_cost": total_cost,
                "num_resources": len(resources)
            }
            
            logger.info(
                f"Processed Excel BOM: {file_path} "
                f"({len(resources)} resources, ${total_cost:.2f}/month)"
            )
            return result
            
        except Exception as e:
            logger.error(f"Failed to process Excel BOM {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def process_document(file_path: str) -> Dict[str, Any]:
        """
        Process any supported document type.
        
        Args:
            file_path: Path to document
            
        Returns:
            Processed document data
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return DocumentProcessor.process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return DocumentProcessor.process_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return DocumentProcessor.process_excel_bom(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    @staticmethod
    def process_multiple_documents(file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of processed documents
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = DocumentProcessor.process_document(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                results.append({
                    "file_path": file_path,
                    "error": str(e),
                    "processed": False
                })
        
        return results


def extract_evidence_from_documents(
    documents: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    Extract migration evidence from processed documents.
    
    Args:
        documents: List of processed documents
        
    Returns:
        Consolidated evidence dictionary
    """
    evidence = {
        "services": [],
        "network": {},
        "compute": [],
        "storage": [],
        "security": {},
        "cost_data": {},
        "diagrams": [],
        "technical_specs": []
    }
    
    for doc in documents:
        if doc.get("file_type") == "excel_bom":
            # Extract cost and resource data from BOM
            evidence["cost_data"]["monthly_total"] = doc.get("total_monthly_cost", 0)
            evidence["cost_data"]["resources"] = doc.get("resources", [])
        
        elif doc.get("file_type") in ["pdf", "docx"]:
            # Extract text evidence
            full_text = doc.get("full_text", "")
            
            # Look for service mentions (simplified keyword search)
            cloud_services = [
                "EC2", "S3", "RDS", "Lambda", "ECS", "EKS",
                "Azure VM", "Blob Storage", "SQL Database",
                "GCE", "Cloud Storage", "Cloud SQL"
            ]
            
            for service in cloud_services:
                if service.lower() in full_text.lower():
                    evidence["services"].append({
                        "service": service,
                        "source": doc["file_path"]
                    })
    
    return evidence
